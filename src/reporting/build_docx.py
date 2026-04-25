"""Genera outputs/suratech_informe_ejecutivo.docx con branding SURA."""

from __future__ import annotations
import json
from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

from src.config import DOCS_DATA, OUTPUTS_DIR

SURA_AZUL = RGBColor(0x00, 0x35, 0x9C)
SURA_AQUA = RGBColor(0x00, 0xAE, 0xC7)
SURA_GRIS_OSCURO = RGBColor(0x1A, 0x1A, 0x1A)
SURA_GRIS_MEDIO  = RGBColor(0x6B, 0x72, 0x80)
ROJO     = RGBColor(0xDC, 0x26, 0x26)
AMARILLO = RGBColor(0xF5, 0x9E, 0x0B)
VERDE    = RGBColor(0x10, 0xB9, 0x81)


def _set_cell_bg(cell, hex_color: str) -> None:
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color)
    tcPr.append(shd)


def _h1(doc: Document, text: str) -> None:
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.font.size = Pt(20); r.font.bold = True; r.font.color.rgb = SURA_AZUL


def _h2(doc: Document, text: str) -> None:
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.font.size = Pt(14); r.font.bold = True; r.font.color.rgb = SURA_AZUL
    p.paragraph_format.space_before = Pt(14)
    p.paragraph_format.space_after = Pt(6)


def _h3(doc: Document, text: str, color: RGBColor = None) -> None:
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.font.size = Pt(11); r.font.bold = True
    r.font.color.rgb = color or SURA_AZUL
    p.paragraph_format.space_before = Pt(8)


def _p(doc: Document, text: str, bold: bool = False, italic: bool = False) -> None:
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.font.size = Pt(11); r.font.bold = bold; r.font.italic = italic
    r.font.color.rgb = SURA_GRIS_OSCURO


def build_docx(data: dict) -> Document:
    doc = Document()

    # Margenes
    for sec in doc.sections:
        sec.top_margin = Cm(2.0); sec.bottom_margin = Cm(2.0)
        sec.left_margin = Cm(2.2); sec.right_margin = Cm(2.2)

    # ===== PORTADA =====
    p = doc.add_paragraph(); r = p.add_run("SURA Tech Colombia")
    r.font.size = Pt(36); r.font.bold = True; r.font.color.rgb = SURA_AZUL
    p = doc.add_paragraph(); r = p.add_run("Informe estratégico mensual")
    r.font.size = Pt(20); r.font.color.rgb = SURA_AQUA
    p = doc.add_paragraph(); r = p.add_run(f"Periodo: {data['meta']['periodo']['mes_actual']}")
    r.font.size = Pt(14); r.font.color.rgb = SURA_GRIS_MEDIO
    doc.add_paragraph()
    p = doc.add_paragraph(); r = p.add_run("Elaborado por Beyond Media Agency")
    r.font.size = Pt(11); r.font.italic = True; r.font.color.rgb = SURA_GRIS_MEDIO
    doc.add_page_break()

    # ===== RESUMEN EJECUTIVO =====
    _h1(doc, "Resumen ejecutivo")
    k = data["kpis_globales"]
    transcurrido = data["meta"]["periodo"].get("porcentaje_transcurrido", 0.7)

    _p(doc, f"Inversión total: ${k['inversion_total']:,.2f} USD")
    _p(doc, f"Leads CRM totales: {k['leads_crm_totales']:,} de {k['compromiso_leads_total']:,} comprometidos")
    cumpl = k['cumplimiento_global_pct']
    txt = f"Cumplimiento global: {cumpl*100:.1f}% (vs {transcurrido*100:.0f}% transcurrido del mes)"
    p = doc.add_paragraph(); r = p.add_run(txt)
    r.font.size = Pt(11); r.font.bold = True
    r.font.color.rgb = ROJO if cumpl/transcurrido < 0.7 else (AMARILLO if cumpl/transcurrido < 0.9 else VERDE)
    _p(doc, f"CPL Negocio promedio: ${k['cpl_negocio_promedio']:.2f} USD")

    # Alertas
    _h2(doc, "Alertas críticas")
    for a in data.get("alertas_criticas", []):
        _h3(doc, "• " + a["titulo"], color=ROJO)
        _p(doc, a["detalle"])
        _p(doc, "Acción: " + a["accion_sugerida"], italic=True)

    doc.add_page_break()

    # ===== POR SEGURO =====
    _h1(doc, "Estado por seguro")
    for s in data["seguros"]:
        c = s["crm"]; ga = s["google_ads"]["kpis"]
        _h2(doc, s["nombre"] + ("" if s["es_principal"] else "  (extra · solo ads)"))
        cump = c.get("cumplimiento_pct")
        if cump is not None:
            ratio = cump/transcurrido
            color = ROJO if ratio < 0.6 else (AMARILLO if ratio < 0.9 else VERDE)
            label = "CRÍTICO" if ratio < 0.6 else ("EN RIESGO" if ratio < 0.9 else "EN OBJETIVO")
            p = doc.add_paragraph(); r = p.add_run(f"  {label} · Cumplimiento {cump*100:.1f}%")
            r.font.bold = True; r.font.color.rgb = color; r.font.size = Pt(12)
            _p(doc, f"  Leads CRM: {c.get('leads_crm'):,} / Compromiso: {c.get('compromiso_leads'):,}")
            _p(doc, f"  CPL Negocio: ${c.get('cpl_negocio') or 0:.2f}")
        else:
            _p(doc, "  Sin tracking CRM. Solo se mide actividad de Google Ads.", italic=True)
        _p(doc, f"  Google Ads: ${ga.get('coste_usd', 0):,.2f} USD invertido · "
                f"{int(ga.get('conversiones') or 0):,} conversiones · "
                f"{s['google_ads']['campanias_count']} campañas activas")

    doc.add_page_break()

    # ===== TOP 10 RECOMENDACIONES =====
    _h1(doc, "TOP 10 — Recomendaciones accionables")
    _p(doc, "Acciones priorizadas para cumplir los compromisos de leads del mes y bajar el CPL Negocio.", italic=True)
    for i, r in enumerate(data.get("recomendaciones", []), 1):
        prio_color = ROJO if r["prioridad"] == "critica" else (AMARILLO if r["prioridad"] == "alta" else VERDE)
        p = doc.add_paragraph()
        run1 = p.add_run(f"#{i}  ")
        run1.font.size = Pt(13); run1.font.bold = True; run1.font.color.rgb = SURA_AZUL
        run2 = p.add_run(f"[{r['prioridad'].upper()}] ")
        run2.font.size = Pt(10); run2.font.bold = True; run2.font.color.rgb = prio_color
        run3 = p.add_run(r["titulo"])
        run3.font.size = Pt(13); run3.font.bold = True; run3.font.color.rgb = SURA_GRIS_OSCURO

        _p(doc, f"   Seguro: {r['seguro']} · Plataforma: {r['plataforma']} · Esfuerzo: {r['esfuerzo']} · Plazo: {r['plazo']}", italic=True)
        _p(doc, f"   ¿Qué hacer? {r['que_hacer']}")
        _p(doc, f"   ¿Por qué? {r['por_que']}")
        _p(doc, f"   Impacto: {r['impacto']}")
        doc.add_paragraph()

    doc.add_page_break()

    # ===== DISCREPANCIAS =====
    _h1(doc, "Discrepancias y gaps")
    d = data.get("discrepancias_y_gaps", {})
    for k_, v_ in d.items():
        _h3(doc, "• " + k_.replace("_", " "))
        _p(doc, v_)

    # ===== CASO DE NEGOCIO =====
    _h1(doc, "Caso de negocio")
    c = data.get("caso_negocio", {})
    _h2(doc, "Estado actual")
    _p(doc, c.get("estado_actual") or "—")
    _h2(doc, "Proyección si aplicamos las top 5 recomendaciones")
    _p(doc, c.get("proyeccion_si_aplicamos_top5") or "—")
    _h2(doc, "Próximo paso")
    _p(doc, c.get("proximo_paso") or "—")

    # Footer
    doc.add_paragraph()
    p = doc.add_paragraph(); r = p.add_run(f"Beyond Media Agency · {data['meta']['generated_at']}")
    r.font.size = Pt(9); r.font.italic = True; r.font.color.rgb = SURA_GRIS_MEDIO
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    return doc


def run() -> int:
    data = json.loads(DOCS_DATA.read_text(encoding="utf-8"))
    doc = build_docx(data)
    OUTPUTS_DIR.mkdir(parents=True, exist_ok=True)
    out = OUTPUTS_DIR / "suratech_informe_ejecutivo.docx"
    doc.save(out)
    print(f"[docx] -> {out}")
    return 0


if __name__ == "__main__":
    raise SystemExit(run())
